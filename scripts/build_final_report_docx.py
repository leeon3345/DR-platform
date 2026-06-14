from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parents[1] / "deliverables"
OUT_PATH = OUT_DIR / "DR-Platform_v1_Final_Report.docx"

FONT_KR = "Apple SD Gothic Neo"
FONT_EN = "Arial"
NAVY = "37352F"
BLUE = "E9E9E7"
HEADER_BLUE = "F7F6F3"
LIGHT_GRAY = "FBFBFA"
MID_GRAY = "E9E9E7"
DARK = "37352F"
MUTED = "787774"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="C9C9C9", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_run_font(run, size=None, color=None, bold=None, italic=None, name=FONT_KR):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_border_bottom(paragraph, color=BLUE, size="8", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def paragraph_border_left(paragraph, color=BLUE, size="20", space="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for node in (fld_begin, instr, fld_sep, fld_text, fld_end):
        run._r.append(node)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_KR
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KR)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = FONT_KR
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KR)
        style.font.bold = True

    styles["Heading 1"].font.size = Pt(22)
    styles["Heading 1"].font.color.rgb = RGBColor.from_string(NAVY)
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 2"].font.size = Pt(13.5)
    styles["Heading 2"].font.color.rgb = RGBColor.from_string(NAVY)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = RGBColor.from_string(NAVY)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)


def setup_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.35)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("DR-Platform Final Report  |  ")
    set_run_font(r, size=9, color=MUTED)
    add_page_number(p)
    for run in p.runs:
        set_run_font(run, size=9, color=MUTED)


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("DR-Platform v1")
    set_run_font(r, size=36, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Kubernetes 다중 클러스터 재해복구(DR) 및 Failback 자동화 플랫폼")
    set_run_font(r, size=12.5, color=MUTED)

    rule_top = doc.add_paragraph()
    rule_top.paragraph_format.space_before = Pt(0)
    rule_top.paragraph_format.space_after = Pt(7)
    paragraph_border_bottom(rule_top, BLUE, size="5", space="1")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("최종 보고서  |  Final Project Report")
    set_run_font(r, size=10, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("1조")
    set_run_font(r, size=10, color=MUTED, bold=True)

    meta = [
        ("과목명", "클라우드 가상화기술"),
        ("팀명", "1조"),
        ("프로젝트명", "DR-Platform"),
        ("문서 버전", "v3.1 (최종보고서)"),
        ("대상 브랜치", "main"),
        ("작성일", "2026. 06. 14"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, width_dxa=7660, indent_dxa=120)
    widths = [1600, 6060]
    for row_idx, (label, value) in enumerate(meta):
        row = table.rows[row_idx]
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_borders(cell, MID_GRAY)
            set_cell_margins(cell, top=70, bottom=70, start=120, end=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(row.cells[0], HEADER_BLUE)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = row.cells[0].paragraphs[0].add_run(label)
        set_run_font(run, size=9.5, color=MUTED, bold=True)
        run = row.cells[1].paragraphs[0].add_run(value)
        set_run_font(run, size=9.5, color=DARK)

    doc.add_page_break()


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(8)
    paragraph_border_bottom(p, BLUE, size="5", space="4")
    run = p.add_run(text)
    set_run_font(run, size=22, color=NAVY, bold=True)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    paragraph_border_left(p, BLUE, size="18", space="5")
    run = p.add_run(text)
    set_run_font(run, size=13.5, color=NAVY, bold=True)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=11, color=DARK, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, size=11, color=DARK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=DARK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=10.7, color=DARK)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=10.7, color=DARK)


def add_table(doc, headers, rows, widths=None, font_size=9.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, width_dxa=9360, indent_dxa=120)
    if widths is None:
        widths = [9360 // len(headers)] * len(headers)
    hdr = table.rows[0]
    for idx, title in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, HEADER_BLUE)
        set_cell_borders(cell, MID_GRAY)
        set_cell_margins(cell, top=95, bottom=95, start=120, end=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        set_run_font(run, size=font_size, color=DARK, bold=True)
    for row_idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cell = cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_borders(cell, MID_GRAY)
            set_cell_margins(cell, top=90, bottom=90, start=120, end=120)
            if row_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(headers) <= 4 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_image_placeholder(doc, caption, height_in=2.8):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, width_dxa=9360, indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_width(cell, 9360)
    set_cell_borders(cell, "D3D3D1", size="8")
    set_cell_shading(cell, "FAFAFA")
    set_cell_margins(cell, top=80, bottom=80, start=120, end=120)
    row = table.rows[0]
    row.height = Inches(height_in)
    row.height_rule = 2
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell.paragraphs[0].add_run("")
    set_run_font(run, size=10, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(caption)
    set_run_font(run, size=10, color=DARK)


def add_toc(doc):
    add_h1(doc, "목차")
    toc = [
        "Executive Summary",
        "1. 보고서 개요",
        "2. 개발 로드맵 진척도 평가",
        "3. 중간보고서 대비 주요 변경점",
        "4. 구현 아키텍처",
        "5. 핵심 구현 기능",
        "6. End-to-End 배포 검증 결과",
        "7. 기능 목록 최종 현황 (MoSCoW 대비)",
        "8. 트러블 슈팅 및 문제 해결",
        "9. 한계 및 향후 과제",
        "10. 결론",
    ]
    for item in toc:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, size=11, color=DARK)
    doc.add_page_break()


def build_doc():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    setup_section(doc.sections[0])
    configure_styles(doc)
    add_cover(doc)

    add_h1(doc, "Executive Summary")
    add_h2(doc, "문제 (Problem)")
    add_body(doc, "Kubernetes 기반 서비스는 클러스터 장애, 네트워크 단절, 스토리지 장애가 발생했을 때 백업·복원 절차가 여러 도구와 권한 경계에 흩어져 있다. 운영자는 장애 대피(Failover)와 원상 복구(Failback)를 같은 흐름에서 추적하기 어렵고, 특히 메인망과 임시망이 분리된 환경에서는 플랫폼이 모든 클러스터에 직접 접근하기 어렵다.")
    add_h2(doc, "해결 방식 (Solution)")
    add_body(doc, "DR-Platform은 Cloud K8s와 Edge K3s를 하나의 재해복구 흐름으로 연결한다. dr-agent가 각 클러스터의 상태를 Heartbeat Push 방식으로 전달하고, Velero/MinIO 기반 백업 저장소와 복구 정책 엔진을 통해 Failover 및 Failback 절차를 자동화하였다. 단, 메인망 복원은 보안과 권한 격리를 위해 관리자가 다운로드한 스크립트를 직접 실행하는 구조로 조정하였다.")
    add_h2(doc, "결과 (Result)")
    add_body(doc, "중간보고서의 Phase 1 및 Phase 1.5 범위는 구현 완료되었고, Phase 2 항목 중 외부 공개 URL(zrok) 기능은 앞당겨 반영되었다. 실제 Node.js 백엔드와 Postgres DB를 배포하여 Cloud K8s 장애 발생, Edge K3s 대피 운영, K3s 최신 데이터 백업, K8s 원상 복구 스크립트 실행까지 End-to-End 시나리오를 검증하였다.")
    add_h2(doc, "진척도 요약")
    add_table(
        doc,
        ["페이즈", "계획 범위", "달성 상태", "비고"],
        [
            ["Phase 1 (MVP / P0)", "다중 클러스터 모니터링, 수동 백업/복원, 대시보드 UI", "완료", "dr-agent 및 관제 대시보드 연동 완료"],
            ["Phase 1.5 (P1)", "Failback 자동화 로직, 복구 정책(Policy), drctl CLI", "완료", "8단계 온보딩 및 복구 스크립트 생성 로직 구현"],
            ["Phase 2 (고도화)", "터널링(zrok) 외부 노출, 스토리지 자동 프로비저닝", "부분", "공개 URL 구현 완료, 스토리지 자동화는 이연"],
        ],
        widths=[1900, 3450, 1250, 2760],
        font_size=8.8,
    )
    doc.add_page_break()
    add_toc(doc)

    add_h1(doc, "1. 보고서 개요")
    add_h2(doc, "1.1 목적")
    add_body(doc, "본 최종보고서는 중간보고서에서 설계한 DR-Platform 아키텍처와 개발 로드맵이 실제 코드로 어떻게 구현되었는지를 검증하고, 설계 대비 변경된 사항을 정리한다. 중간보고서가 “무엇을 만들 것인가(설계)”였다면, 본 보고서는 “무엇을 만들었는가(구현 결과)”와 “왜 설계와 달라졌는가(변경 근거)”를 시스템 및 코드 레벨에서 기술한다.")
    add_body(doc, "검증 핵심 산출물은 Cloud K8s(메인망) 및 Edge K3s(임시망) 멀티 클러스터 환경을 구축하고, 실제 서비스 장애 시뮬레이션을 통해 Failover(장애 대피)부터 Failback(원상 복구)까지의 End-to-End 파이프라인이 정상 동작함을 확인한 것이다.")
    add_h2(doc, "1.2 전체 진척도 한눈에 보기")
    add_table(
        doc,
        ["페이즈", "계획 범위", "달성 상태", "비고"],
        [
            ["Phase 1 (MVP / P0)", "다중 클러스터 모니터링, 수동 백업/복원, 대시보드 UI", "완료", "dr-agent 및 관제 대시보드 연동 완료"],
            ["Phase 1.5 (P1)", "Failback 자동화 로직, 복구 정책(Policy), drctl CLI", "완료", "8단계 온보딩 및 복구 스크립트 생성 로직 구현"],
            ["Phase 2 (고도화)", "터널링(zrok) 외부 노출, 스토리지 자동 프로비저닝", "부분", "터널링 및 공개 URL 구현 완료, 스토리지는 이연"],
        ],
        widths=[1900, 3450, 1250, 2760],
        font_size=8.8,
    )

    add_h1(doc, "2. 개발 로드맵 진척도 평가")
    add_h2(doc, "2.1 Phase 1 (MVP) — 완료")
    add_body(doc, "마일스톤: “다중 클러스터 상태 수집 및 Velero 백업 연동”")
    for item in [
        "dr-agent 구현: 대상 클러스터에 배포되어 워크로드 상태 및 백업 현황을 수집하는 에이전트(Heartbeat 방식).",
        "중앙 관제 대시보드: React 기반 UI를 통해 여러 클러스터의 상태 및 경고(Alert)를 시각화.",
        "백업/복원 기초 구현: Velero를 이용한 네임스페이스 단위의 K8s 리소스 백업 및 S3(MinIO) 업로드 기능.",
    ]:
        add_bullet(doc, item)
    add_h2(doc, "2.2 Phase 1.5 (P1) — 완료")
    add_body(doc, "마일스톤: “Failback 파이프라인 자동화 및 CLI 도구 배포”")
    for item in [
        "Failover & Failback 엔진: 장애 발생 시 Edge K3s로 복원하고, 복구 완료 후 K3s의 최신 데이터를 다시 K8s로 돌려보내는 자동화 파이프라인.",
        "복구 스크립트 발급기: Failback 진행 시 플랫폼이 자동 백업을 수행하고 사용자가 실행할 수 있는 restore-*.sh 다운로드 기능 구현.",
        "drctl CLI: 터미널에서 플랫폼 연동(init) 및 복구 정책(policy set)을 손쉽게 설정할 수 있는 CLI 도구 배포.",
    ]:
        add_bullet(doc, item)
    add_h2(doc, "2.3 Phase 2 (고도화) — 부분 구현")
    add_body(doc, "Phase 2 과제 중 외부망 접근 제약 해소를 위한 zrok 기반 터널링과 공개 URL 지원을 앞당겨 구현하였다. 반면 클라우드 네이티브 스토리지(EBS, EFS 등) 자동 프로비저닝은 로컬 MinIO 기반의 인프라 제약으로 인해 향후 과제로 이연하였다.")

    add_h1(doc, "3. 중간보고서 대비 주요 변경점")
    add_table(
        doc,
        ["#", "구분", "중간보고서(설계)", "최종(구현)", "변경 근거"],
        [
            ["1", "Failback 실행", "플랫폼이 타겟 클러스터에 자동 복원 명령", "사용자 수동 스크립트 실행", "K8s 메인망에 데이터를 덮어쓰는 작업의 치명성을 고려하여 망 분리 및 권한 격리 원칙에 맞게 변경."],
            ["2", "상태 동기화", "플랫폼이 클러스터를 Polling", "Agent Heartbeat Push", "사설망 클러스터 접근 불가 문제를 해결하기 위해 에이전트가 중앙 서버로 상태를 Push."],
            ["3", "스토리지 격리", "네임스페이스 기반 논리 격리", "Velero Prefix 기반 격리", "단일 MinIO 공유 환경에서 prefix=<CLUSTER_ID>/ 구조로 사용자별 백업 데이터를 분리."],
            ["4", "UI/UX 온보딩", "통합 스크립트 제공", "8단계 상세 가이드라인", "데모 투명성 및 기술적 신뢰도 확보를 위해 내부 과정을 단계별로 노출."],
            ["5", "모니터링 스택", "Zabbix", "Prometheus", "K8s 환경에 더 최적화된 성능과 연동성을 제공하고 설정 복잡도를 낮춤."],
        ],
        widths=[500, 1400, 2100, 1900, 3460],
        font_size=8.0,
    )

    add_h1(doc, "4. 구현 아키텍처")
    add_h2(doc, "4.1 4계층 구조")
    for item in [
        "Presentation (React 18): 클러스터 헬스 체크, Failback 마법사, 알림 및 토폴로지 시각화 패널 제공.",
        "Business Logic (Node.js/Express): Agent Heartbeat 수신 엔진, 복구 정책 관리, Firing/Resolved 알람 상태 전이 엔진(server.mjs).",
        "Integration (Velero & CLI): S3(MinIO) 스토리지 연동, 백업/복원 REST API 연동, drctl을 통한 설정 자동화.",
        "K8s Cluster: Cloud K8s(Primary / 메인망)와 Edge K3s(Recovery / 임시망)로 구성.",
    ]:
        add_bullet(doc, item)
    add_image_placeholder(doc, "그림 1. DR-Platform 구현 아키텍처")
    add_h2(doc, "4.2 전역 상태 관리 및 이벤트 엔진")
    add_body(doc, "Heartbeat 기반 알람 해소 로직: K3s 클러스터 복구 완료 후, 정상 동작 중인 네임스페이스 정보를 수신하면 과거 메인망에서 발생했던 Firing 상태의 알람을 Resolved로 자동 전환하는 이벤트 엔진을 server.mjs에 구현하였다.")
    add_body(doc, "비동기 Operation 큐: 백업/복원과 같이 장시간이 소요되는 작업은 Operation 큐로 관리하고 프론트엔드는 이를 폴링하여 진행률(%)을 렌더링한다.")

    add_h1(doc, "5. 핵심 구현 기능")
    feature_sections = [
        ("5.1 Heartbeat 기반 다중 클러스터 동기화 (dr-agent)", [
            "통신 아키텍처: 플랫폼 서버가 사설망에 갇힌 클러스터에 직접 접근할 수 없는 네트워크 구조적 한계를 극복하기 위해, 에이전트가 중앙 서버로 주도적으로 데이터를 보내는 Heartbeat Push 모델을 채택하였다.",
            "지표 수집: Helm으로 배포된 dr-agent는 30초 단위로 클러스터의 노드 상태, 파드 레플리카(Ready/Total), Velero 백업 및 복원 이력(Phase, Expiration)을 server.mjs의 receiveAgentHeartbeat API로 전송한다.",
            "FSM 기반 알람 자동 해소: 장애 발생 시 생성된 Firing 상태의 알림은 클러스터가 복구되어 워크로드가 Running 상태로 보고되면 서버 측 이벤트 엔진이 이를 인지하여 Resolved 상태로 전이한다.",
        ]),
        ("5.2 Policy-as-Code 기반 SLA 관리 엔진", [
            "CLI 기반 선언적 관리: drctl policy set 명령어를 통해 각 네임스페이스별 RTO, RPO, 서비스 등급(Tier: Critical/Standard/Low)을 선언적으로 정의한다.",
            "Backup Freshness 모니터링: 대시보드는 등록된 Policy(RPO)와 Velero가 보고한 가장 최근 성공 백업 완료 시간을 비교하여 RPO 초과 시 UI 경고를 노출한다.",
        ]),
        ("5.3 Failback 자동화 파이프라인 및 복구 스크립트 발급", [
            "양방향 Failback 오케스트레이션: 대시보드에서 Failback을 트리거하면 서버가 K3s 에이전트의 큐에 BackupCommand를 삽입하고, 에이전트가 K3s 상에서 Velero Backup을 자동 실행한다.",
            "비동기 상태 폴링: 프론트엔드는 failback 엔드포인트를 폴링하며 Queued → Submitted → InProgress → Completed 전이와 진행률을 실시간 렌더링한다.",
            "권한 분리형 스크립트 제공: 백업 완료 후 restore-<backupName>.sh를 생성하고, 관리자가 메인망에서 직접 실행하도록 하여 보안 및 권한 격리 원칙을 준수하였다.",
        ]),
        ("5.4 Velero Prefix 기반 멀티테넌시 및 데이터 격리", [
            "스토리지 통합 및 격리: 여러 테넌트가 백업 저장소로 단일 S3(MinIO) 버킷을 공유하는 환경을 구성하였다.",
            "디렉토리 논리적 격리: velero install 단계에서 --backup-location-config prefix=<CLUSTER_ID>/ 옵션을 강제하여 각 클러스터 백업 데이터가 고유 디렉토리에 저장되도록 설계하였다.",
            "보안성 확보: A 클러스터의 Velero는 A/ 디렉토리 내의 백업 객체만 인지하므로 다른 사용자의 데이터를 복원하거나 덮어쓰는 논리적 데이터 오염 사고를 차단한다.",
        ]),
    ]
    for title, items in feature_sections:
        add_h2(doc, title)
        for item in items:
            add_bullet(doc, item)

    add_h1(doc, "6. End-to-End 배포 검증 결과")
    add_body(doc, "실제 Node.js 백엔드 앱과 Postgres DB를 K8s에 배포하여 재해복구 시나리오를 검증하였다.")
    steps = [
        "배포 전 초기 상태: Cloud K8s에서 서비스 정상 동작 및 주기적 백업 진행.",
        "Stage 1 (장애 발생 및 Failover): Cloud K8s의 장애를 시뮬레이션하고, 대피소인 Edge K3s에서 이전 백업본을 가져와 시스템 복원.",
        "Stage 2 (K3s 임시 운영): K3s 환경에서 서비스가 구동되며 새로운 게시글 및 DB 데이터가 누적됨.",
        "Stage 3 (Failback 버튼 클릭): 대시보드에서 Fail back to primary cluster를 실행하고 플랫폼이 K3s의 최신 상태를 MinIO에 백업 완료.",
        "최종 검증: 다운로드받은 .sh 스크립트를 복구된 K8s 터미널에서 실행하여 K3s에서 생성했던 최신 게시글 데이터까지 유실 없이 원상 복구됨을 확인.",
    ]
    for step in steps:
        add_numbered(doc, step)
    add_image_placeholder(doc, "그림 2. Failover 및 Failback 실배포 검증 화면")

    add_h1(doc, "7. 기능 목록 최종 현황 (MoSCoW 대비)")
    add_table(
        doc,
        ["ID", "기능명", "MoSCoW", "상태"],
        [
            ["F-01", "drctl CLI 연동 도구", "Must", "완료"],
            ["F-02", "dr-agent Helm Chart 패키징", "Must", "완료"],
            ["F-03", "대시보드 상태 모니터링 패널", "Must", "완료"],
            ["F-04", "복구 정책(Policy-as-Code) 관리", "Must", "완료"],
            ["F-05", "Failover(K3s 대피) 파이프라인", "Must", "완료"],
            ["F-06", "Failback 스크립트 자동 생성기", "Should", "완료"],
            ["F-07", "멀티테넌시 스토리지 버킷 격리", "Should", "완료"],
            ["F-08", "8단계 온보딩 UI 마법사", "Should", "완료"],
            ["F-09", "터널링(zrok) 기반 공개 URL 제공", "Could", "완료 (앞당김)"],
            ["F-10", "타겟 클러스터 자동 Failback", "Could", "보류 (수동 스크립트로 대체)"],
            ["F-11", "클라우드 스토리지 자동 프로비저닝", "Could", "이연 (Phase 2)"],
        ],
        widths=[900, 4200, 1500, 2760],
        font_size=8.8,
    )
    add_body(doc, "Must·Should 기능은 전량 완료되었으며, Could에 해당하던 외부 URL 노출 기능을 앞당겨 구현하였다. 치명적 보안 위협을 방지하기 위해 완전 자동 Failback(F-10) 설계는 보류하고 수동 스크립트 발급 구조로 변경하였다.")

    add_h1(doc, "8. 트러블 슈팅 및 문제 해결")
    issues = [
        ("8.1 관리형 클러스터 권한 누락 및 UI 블랭크 이슈", [
            "증상: 사용자 토큰에서 중앙 클러스터(cloud-primary, edge-recovery) 접근 권한이 누락되면서 대시보드 렌더링에 필요한 필수 상태를 가져오지 못해 UI가 깨지거나 멈춤.",
            "원인: UI 컴포넌트가 중앙 클러스터 상태에 의존했으나, 백엔드 API가 접근을 엄격하게 404 에러로 차단.",
            "해결: tokens.json에 필수 클러스터 권한을 복구하고, App.jsx에서 중앙 클러스터가 좌측 사이드바에 노출되지 않도록 필터링 로직을 추가하여 권한과 UI 가시성을 분리.",
        ]),
        ("8.2 가상(Mock) 환경 SSH 실행 오류로 인한 API Error 도배", [
            "증상: UI 상단 메트릭 패널 전체가 빨간색 API error로 도배되고 상단 타이틀이 cloud-primary로 고정됨.",
            "원인: 프론트엔드 useEffect 훅이 초기 렌더링 시 사용자의 test 클러스터 대신 첫 번째 항목인 cloud-primary를 강제로 선택.",
            "해결: 활성화된 클러스터가 없을 경우 중앙 관리형 클러스터를 건너뛰고 사용자의 실제 클러스터(test)를 우선 자동 선택하도록 수정.",
        ]),
        ("8.3 Helm Repository 업데이트 JSON Unmarshaling 에러", [
            "증상: helm repo update 실행 시 cannot unmarshal string into Go value of type repo.IndexFile 에러 발생.",
            "원인: 외부 터널 주소가 백엔드 API 서버가 아닌 Vite 개발 서버를 직접 가리켜 Helm이 JSON이 아닌 HTML 응답을 수신.",
            "해결: vite.config.js에 /api 요청을 백엔드 서버(127.0.0.1:3001)로 포워딩하는 프록시 설정을 추가.",
        ]),
        ("8.4 Agent 설치 후 대시보드 맵에서 클러스터 에러 지속", [
            "증상: DR Agent를 정상적으로 Helm 설치했음에도 대시보드에서 녹색(Safe)이 아니라 계속 에러 상태로 노출.",
            "원인: 가이드의 가짜 토큰(my-platform-token)이 그대로 지정되어 백엔드가 Heartbeat를 차단했고, Proxy 연동 누락으로 외부 터널 신호가 백엔드까지 도달하지 못함.",
            "해결: 사용자 클러스터에 할당된 실제 토큰과 네임스페이스 옵션을 추가하여 Helm 배포를 재진행하고, 프론트-백엔드 Proxy 서버를 재시작.",
        ]),
    ]
    for title, bullets in issues:
        add_h2(doc, title)
        for item in bullets:
            prefix = item.split(":")[0] + ":"
            add_body(doc, item, bold_prefix=prefix)

    add_h1(doc, "9. 한계 및 향후 과제")
    for item in [
        "완전 자동화된 Failback의 부재: 현재 아키텍처는 권한 제약을 피하기 위해 사용자가 직접 메인망에서 스크립트를 실행해야 한다. 향후 ArgoCD 등을 활용한 GitOps 방식과 연계하면 수동 개입 없는 자동 복구 파이프라인으로 고도화할 수 있다.",
        "스토리지 벤더 종속성 최소화: 현재 테스트는 S3 호환 로컬 스토리지인 MinIO 환경에 맞춰져 있다. 프로덕션 배포를 위해 AWS EBS, EFS 및 GCP 영구 디스크(PD)의 볼륨 스냅샷 기능(CSI 훅)까지 포괄하는 확장이 필요하다.",
        "원클릭 설치 스크립트 지원: 투명성을 위한 8단계 설치 과정이 온보딩 장벽이 될 수 있다. 향후 curl 기반 원클릭 스크립트 버전을 추가로 지원할 계획이다.",
    ]:
        add_numbered(doc, item)

    add_h1(doc, "10. 결론")
    add_body(doc, "DR-Platform은 중간보고서에서 설계한 K8s 다중 클러스터 재해복구 파이프라인을 성공적으로 구현하였다. 핵심 마일스톤인 “장애 발생(K8s) → 대피소 복원(K3s) → 최신 데이터 백업 → 원상 복구 스크립트(.sh) 실행(K8s)”으로 이어지는 End-to-End 재해복구 사이클이 검증되었다.")
    add_body(doc, "설계 대비 일부 변경(스크립트 발급 방식, Heartbeat 도입 등)은 모두 시스템 보안 및 물리적 망 분리 환경을 고려한 현실적 판단에 따른 것이다. 결과적으로 플랫폼의 안정성과 데이터 격리성을 강화하였으며, 이연된 스토리지 프로비저닝 및 GitOps 연계 과제는 상용화를 위한 향후 고도화 목표로 남는다.")

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_doc()
